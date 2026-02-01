"""Document parser for PDF, DOCX, and TXT brand documents.

Parses brand documents uploaded by users and extracts text content
for further processing by the brand extraction service.

Supported formats:
- PDF: Uses pypdf for extraction
- DOCX: Uses python-docx for extraction
- TXT: Plain text reading with encoding detection

ERROR LOGGING REQUIREMENTS:
- Log method entry/exit at DEBUG level with parameters (sanitized)
- Log all exceptions with full stack trace and context
- Include entity IDs (project_id, page_id) in all service logs
- Log validation failures with field names and rejected values
- Log state transitions (phase changes) at INFO level
- Add timing logs for operations >1 second
"""

import io
import time
import traceback
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import BinaryIO

from app.core.logging import get_logger

logger = get_logger(__name__)

# Constants
SLOW_OPERATION_THRESHOLD_MS = 1000
MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50MB
MAX_FILENAME_LOG_LENGTH = 50
MAX_CONTENT_PREVIEW_LENGTH = 200

# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class DocumentFormat(Enum):
    """Supported document formats."""

    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


@dataclass
class DocumentMetadata:
    """Metadata extracted from a document."""

    filename: str
    format: DocumentFormat
    file_size_bytes: int
    page_count: int | None = None  # Only for PDF
    word_count: int = 0
    character_count: int = 0
    author: str | None = None
    title: str | None = None
    created_date: str | None = None
    modified_date: str | None = None

    def to_dict(self) -> dict[str, str | int | None]:
        """Convert to dictionary for serialization."""
        return {
            "filename": self.filename,
            "format": self.format.value,
            "file_size_bytes": self.file_size_bytes,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "character_count": self.character_count,
            "author": self.author,
            "title": self.title,
            "created_date": self.created_date,
            "modified_date": self.modified_date,
        }


@dataclass
class DocumentParseResult:
    """Result of document parsing operation."""

    success: bool
    content: str = ""
    metadata: DocumentMetadata | None = None
    error: str | None = None
    duration_ms: float = 0.0
    sections: list[str] = field(default_factory=list)

    @property
    def is_empty(self) -> bool:
        """Check if parsed content is empty."""
        return not self.content or not self.content.strip()


class DocumentParserError(Exception):
    """Base exception for document parser errors."""

    def __init__(
        self,
        message: str,
        project_id: str | None = None,
        filename: str | None = None,
    ) -> None:
        super().__init__(message)
        self.project_id = project_id
        self.filename = filename


class UnsupportedFormatError(DocumentParserError):
    """Raised when document format is not supported."""

    def __init__(
        self,
        extension: str,
        project_id: str | None = None,
        filename: str | None = None,
    ) -> None:
        super().__init__(
            f"Unsupported document format: {extension}. "
            f"Supported formats: {', '.join(SUPPORTED_EXTENSIONS)}",
            project_id=project_id,
            filename=filename,
        )
        self.extension = extension


class FileTooLargeError(DocumentParserError):
    """Raised when file exceeds maximum size limit."""

    def __init__(
        self,
        file_size: int,
        max_size: int,
        project_id: str | None = None,
        filename: str | None = None,
    ) -> None:
        super().__init__(
            f"File size ({file_size:,} bytes) exceeds maximum "
            f"allowed size ({max_size:,} bytes)",
            project_id=project_id,
            filename=filename,
        )
        self.file_size = file_size
        self.max_size = max_size


class DocumentCorruptedError(DocumentParserError):
    """Raised when document is corrupted or unreadable."""

    pass


class DocumentParser:
    """Parser for PDF, DOCX, and TXT brand documents.

    Features:
    - Parses PDF, DOCX, and TXT files
    - Extracts text content and metadata
    - Validates file format and size
    - Comprehensive logging per requirements

    Usage:
        parser = DocumentParser()
        result = parser.parse_file(file_path)
        # or
        result = parser.parse_bytes(file_bytes, filename)
    """

    def __init__(
        self,
        max_file_size: int = MAX_FILE_SIZE_BYTES,
    ) -> None:
        """Initialize document parser.

        Args:
            max_file_size: Maximum allowed file size in bytes.
        """
        self._max_file_size = max_file_size

        logger.debug(
            "DocumentParser initialized",
            extra={
                "max_file_size_bytes": self._max_file_size,
                "supported_formats": list(SUPPORTED_EXTENSIONS),
            },
        )

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for logging (truncate if too long)."""
        if len(filename) > MAX_FILENAME_LOG_LENGTH:
            return filename[:MAX_FILENAME_LOG_LENGTH] + "..."
        return filename

    def _detect_format(self, filename: str) -> DocumentFormat:
        """Detect document format from filename extension.

        Args:
            filename: Name of the file

        Returns:
            Detected document format

        Raises:
            UnsupportedFormatError: If format is not supported
        """
        extension = Path(filename).suffix.lower()

        if extension == ".pdf":
            return DocumentFormat.PDF
        elif extension == ".docx":
            return DocumentFormat.DOCX
        elif extension == ".txt":
            return DocumentFormat.TXT
        else:
            raise UnsupportedFormatError(extension, filename=filename)

    def _validate_file_size(
        self,
        size: int,
        filename: str,
        project_id: str | None = None,
    ) -> None:
        """Validate file size is within limits.

        Args:
            size: File size in bytes
            filename: Name of the file
            project_id: Project ID for logging

        Raises:
            FileTooLargeError: If file exceeds size limit
        """
        if size > self._max_file_size:
            logger.warning(
                "File size validation failed",
                extra={
                    "field": "file_size",
                    "value": size,
                    "max_allowed": self._max_file_size,
                    "document_filename": self._sanitize_filename(filename),
                    "project_id": project_id,
                },
            )
            raise FileTooLargeError(
                size,
                self._max_file_size,
                project_id=project_id,
                filename=filename,
            )

    def _count_words(self, text: str) -> int:
        """Count words in text."""
        return len(text.split())

    def _parse_pdf(
        self,
        file_stream: BinaryIO,
        filename: str,
        file_size: int,
        project_id: str | None = None,
    ) -> DocumentParseResult:
        """Parse PDF document.

        Args:
            file_stream: Binary file stream
            filename: Name of the file
            file_size: File size in bytes
            project_id: Project ID for logging

        Returns:
            DocumentParseResult with extracted content
        """
        try:
            from pypdf import PdfReader
        except ImportError as e:
            logger.error(
                "pypdf not installed",
                extra={
                    "error": str(e),
                    "document_filename": self._sanitize_filename(filename),
                    "project_id": project_id,
                    "stack_trace": traceback.format_exc(),
                },
                exc_info=True,
            )
            raise DocumentParserError(
                "PDF parsing library not available",
                project_id=project_id,
                filename=filename,
            ) from e

        start_time = time.monotonic()

        logger.debug(
            "Parsing PDF document",
            extra={
                "document_filename": self._sanitize_filename(filename),
                "file_size_bytes": file_size,
                "project_id": project_id,
            },
        )

        try:
            reader = PdfReader(file_stream)
            page_count = len(reader.pages)

            # Extract text from all pages
            sections: list[str] = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        sections.append(page_text)
                except Exception as page_error:
                    logger.warning(
                        "Failed to extract text from PDF page",
                        extra={
                            "page_number": i + 1,
                            "document_filename": self._sanitize_filename(filename),
                            "error": str(page_error),
                            "project_id": project_id,
                        },
                    )
                    # Continue with other pages

            content = "\n\n".join(sections)

            # Extract metadata
            pdf_metadata = reader.metadata
            author: str | None = None
            title: str | None = None
            created: str | None = None
            modified: str | None = None

            if pdf_metadata is not None:
                author = str(pdf_metadata.author) if pdf_metadata.author else None
                title = str(pdf_metadata.title) if pdf_metadata.title else None
                created = (
                    str(pdf_metadata.creation_date)
                    if pdf_metadata.creation_date
                    else None
                )
                modified = (
                    str(pdf_metadata.modification_date)
                    if pdf_metadata.modification_date
                    else None
                )

            metadata = DocumentMetadata(
                filename=filename,
                format=DocumentFormat.PDF,
                file_size_bytes=file_size,
                page_count=page_count,
                word_count=self._count_words(content),
                character_count=len(content),
                author=author,
                title=title,
                created_date=created,
                modified_date=modified,
            )

            duration_ms = (time.monotonic() - start_time) * 1000

            logger.debug(
                "PDF parsing complete",
                extra={
                    "document_filename": self._sanitize_filename(filename),
                    "page_count": page_count,
                    "word_count": metadata.word_count,
                    "duration_ms": round(duration_ms, 2),
                    "project_id": project_id,
                },
            )

            if duration_ms > SLOW_OPERATION_THRESHOLD_MS:
                logger.warning(
                    "Slow PDF parsing operation",
                    extra={
                        "document_filename": self._sanitize_filename(filename),
                        "duration_ms": round(duration_ms, 2),
                        "threshold_ms": SLOW_OPERATION_THRESHOLD_MS,
                        "project_id": project_id,
                    },
                )

            return DocumentParseResult(
                success=True,
                content=content,
                metadata=metadata,
                duration_ms=duration_ms,
                sections=sections,
            )

        except Exception as e:
            duration_ms = (time.monotonic() - start_time) * 1000
            logger.error(
                "PDF parsing failed",
                extra={
                    "document_filename": self._sanitize_filename(filename),
                    "error": str(e),
                    "error_type": type(e).__name__,
                    "duration_ms": round(duration_ms, 2),
                    "project_id": project_id,
                    "stack_trace": traceback.format_exc(),
                },
                exc_info=True,
            )
            raise DocumentCorruptedError(
                f"Failed to parse PDF: {e}",
                project_id=project_id,
                filename=filename,
            ) from e

    def _parse_docx(
        self,
        file_stream: BinaryIO,
        filename: str,
        file_size: int,
        project_id: str | None = None,
    ) -> DocumentParseResult:
        """Parse DOCX document.

        Args:
            file_stream: Binary file stream
            filename: Name of the file
            file_size: File size in bytes
            project_id: Project ID for logging

        Returns:
            DocumentParseResult with extracted content
        """
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError as e:
            logger.error(
                "python-docx not installed",
                extra={
                    "error": str(e),
                    "document_filename": self._sanitize_filename(filename),
                    "project_id": project_id,
                    "stack_trace": traceback.format_exc(),
                },
                exc_info=True,
            )
            raise DocumentParserError(
                "DOCX parsing library not available",
                project_id=project_id,
                filename=filename,
            ) from e

        start_time = time.monotonic()

        logger.debug(
            "Parsing DOCX document",
            extra={
                "document_filename": self._sanitize_filename(filename),
                "file_size_bytes": file_size,
                "project_id": project_id,
            },
        )

        try:
            doc = Document(file_stream)

            # Extract text from paragraphs
            sections: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    sections.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        sections.append(row_text)

            content = "\n\n".join(sections)

            # Extract metadata from core properties
            core_props = doc.core_properties
            author = core_props.author if core_props.author else None
            title = core_props.title if core_props.title else None
            created = str(core_props.created) if core_props.created else None
            modified = str(core_props.modified) if core_props.modified else None

            metadata = DocumentMetadata(
                filename=filename,
                format=DocumentFormat.DOCX,
                file_size_bytes=file_size,
                page_count=None,  # DOCX doesn't have reliable page count
                word_count=self._count_words(content),
                character_count=len(content),
                author=author,
                title=title,
                created_date=created,
                modified_date=modified,
            )

            duration_ms = (time.monotonic() - start_time) * 1000

            logger.debug(
                "DOCX parsing complete",
                extra={
                    "document_filename": self._sanitize_filename(filename),
                    "paragraph_count": len(doc.paragraphs),
                    "word_count": metadata.word_count,
                    "duration_ms": round(duration_ms, 2),
                    "project_id": project_id,
                },
            )

            if duration_ms > SLOW_OPERATION_THRESHOLD_MS:
                logger.warning(
                    "Slow DOCX parsing operation",
                    extra={
                        "document_filename": self._sanitize_filename(filename),
                        "duration_ms": round(duration_ms, 2),
                        "threshold_ms": SLOW_OPERATION_THRESHOLD_MS,
                        "project_id": project_id,
                    },
                )

            return DocumentParseResult(
                success=True,
                content=content,
                metadata=metadata,
                duration_ms=duration_ms,
                sections=sections,
            )

        except PackageNotFoundError as e:
            duration_ms = (time.monotonic() - start_time) * 1000
            logger.error(
                "Invalid DOCX file (package not found)",
                extra={
                    "document_filename": self._sanitize_filename(filename),
                    "error": str(e),
                    "error_type": type(e).__name__,
                    "duration_ms": round(duration_ms, 2),
                    "project_id": project_id,
                    "stack_trace": traceback.format_exc(),
                },
                exc_info=True,
            )
            raise DocumentCorruptedError(
                f"Invalid DOCX file: {e}",
                project_id=project_id,
                filename=filename,
            ) from e
        except Exception as e:
            duration_ms = (time.monotonic() - start_time) * 1000
            logger.error(
                "DOCX parsing failed",
                extra={
                    "document_filename": self._sanitize_filename(filename),
                    "error": str(e),
                    "error_type": type(e).__name__,
                    "duration_ms": round(duration_ms, 2),
                    "project_id": project_id,
                    "stack_trace": traceback.format_exc(),
                },
                exc_info=True,
            )
            raise DocumentCorruptedError(
                f"Failed to parse DOCX: {e}",
                project_id=project_id,
                filename=filename,
            ) from e

    def _parse_txt(
        self,
        file_stream: BinaryIO,
        filename: str,
        file_size: int,
        project_id: str | None = None,
    ) -> DocumentParseResult:
        """Parse TXT document.

        Args:
            file_stream: Binary file stream
            filename: Name of the file
            file_size: File size in bytes
            project_id: Project ID for logging

        Returns:
            DocumentParseResult with extracted content
        """
        start_time = time.monotonic()

        logger.debug(
            "Parsing TXT document",
            extra={
                "document_filename": self._sanitize_filename(filename),
                "file_size_bytes": file_size,
                "project_id": project_id,
            },
        )

        try:
            raw_bytes = file_stream.read()

            # Try UTF-8 first, then fallback to other encodings
            encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
            content = ""
            detected_encoding = None

            for encoding in encodings:
                try:
                    content = raw_bytes.decode(encoding)
                    detected_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if not detected_encoding:
                # Last resort: decode with errors='replace'
                content = raw_bytes.decode("utf-8", errors="replace")
                detected_encoding = "utf-8 (with replacements)"
                logger.warning(
                    "TXT encoding detection failed, using fallback",
                    extra={
                        "document_filename": self._sanitize_filename(filename),
                        "fallback_encoding": detected_encoding,
                        "project_id": project_id,
                    },
                )

            # Split into sections by double newlines
            sections = [
                s.strip() for s in content.split("\n\n") if s.strip()
            ]

            metadata = DocumentMetadata(
                filename=filename,
                format=DocumentFormat.TXT,
                file_size_bytes=file_size,
                page_count=None,
                word_count=self._count_words(content),
                character_count=len(content),
            )

            duration_ms = (time.monotonic() - start_time) * 1000

            logger.debug(
                "TXT parsing complete",
                extra={
                    "document_filename": self._sanitize_filename(filename),
                    "encoding": detected_encoding,
                    "word_count": metadata.word_count,
                    "duration_ms": round(duration_ms, 2),
                    "project_id": project_id,
                },
            )

            if duration_ms > SLOW_OPERATION_THRESHOLD_MS:
                logger.warning(
                    "Slow TXT parsing operation",
                    extra={
                        "document_filename": self._sanitize_filename(filename),
                        "duration_ms": round(duration_ms, 2),
                        "threshold_ms": SLOW_OPERATION_THRESHOLD_MS,
                        "project_id": project_id,
                    },
                )

            return DocumentParseResult(
                success=True,
                content=content,
                metadata=metadata,
                duration_ms=duration_ms,
                sections=sections,
            )

        except Exception as e:
            duration_ms = (time.monotonic() - start_time) * 1000
            logger.error(
                "TXT parsing failed",
                extra={
                    "document_filename": self._sanitize_filename(filename),
                    "error": str(e),
                    "error_type": type(e).__name__,
                    "duration_ms": round(duration_ms, 2),
                    "project_id": project_id,
                    "stack_trace": traceback.format_exc(),
                },
                exc_info=True,
            )
            raise DocumentCorruptedError(
                f"Failed to parse TXT: {e}",
                project_id=project_id,
                filename=filename,
            ) from e

    def parse_bytes(
        self,
        file_bytes: bytes,
        filename: str,
        project_id: str | None = None,
    ) -> DocumentParseResult:
        """Parse document from bytes.

        Args:
            file_bytes: Document content as bytes
            filename: Name of the file (used for format detection)
            project_id: Project ID for logging context

        Returns:
            DocumentParseResult with extracted content

        Raises:
            UnsupportedFormatError: If format is not supported
            FileTooLargeError: If file exceeds size limit
            DocumentCorruptedError: If document is corrupted
        """
        start_time = time.monotonic()

        logger.debug(
            "Document parsing started (bytes)",
            extra={
                "document_filename": self._sanitize_filename(filename),
                "file_size_bytes": len(file_bytes),
                "project_id": project_id,
            },
        )

        try:
            # Validate file size
            self._validate_file_size(len(file_bytes), filename, project_id)

            # Detect format
            doc_format = self._detect_format(filename)

            logger.info(
                "Parsing document",
                extra={
                    "document_filename": self._sanitize_filename(filename),
                    "format": doc_format.value,
                    "file_size_bytes": len(file_bytes),
                    "project_id": project_id,
                },
            )

            # Create file-like object from bytes
            file_stream = io.BytesIO(file_bytes)

            # Parse based on format
            if doc_format == DocumentFormat.PDF:
                result = self._parse_pdf(
                    file_stream, filename, len(file_bytes), project_id
                )
            elif doc_format == DocumentFormat.DOCX:
                result = self._parse_docx(
                    file_stream, filename, len(file_bytes), project_id
                )
            else:  # TXT
                result = self._parse_txt(
                    file_stream, filename, len(file_bytes), project_id
                )

            duration_ms = (time.monotonic() - start_time) * 1000

            logger.info(
                "Document parsing complete",
                extra={
                    "document_filename": self._sanitize_filename(filename),
                    "format": doc_format.value,
                    "success": result.success,
                    "word_count": result.metadata.word_count if result.metadata else 0,
                    "duration_ms": round(duration_ms, 2),
                    "project_id": project_id,
                },
            )

            return result

        except (
            UnsupportedFormatError,
            FileTooLargeError,
            DocumentCorruptedError,
            DocumentParserError,
        ):
            # Re-raise known exceptions
            raise
        except Exception as e:
            duration_ms = (time.monotonic() - start_time) * 1000
            logger.error(
                "Document parsing unexpected error",
                extra={
                    "document_filename": self._sanitize_filename(filename),
                    "error": str(e),
                    "error_type": type(e).__name__,
                    "duration_ms": round(duration_ms, 2),
                    "project_id": project_id,
                    "stack_trace": traceback.format_exc(),
                },
                exc_info=True,
            )
            return DocumentParseResult(
                success=False,
                error=f"Unexpected error: {e}",
                duration_ms=duration_ms,
            )

    def parse_file(
        self,
        file_path: str | Path,
        project_id: str | None = None,
    ) -> DocumentParseResult:
        """Parse document from file path.

        Args:
            file_path: Path to the document file
            project_id: Project ID for logging context

        Returns:
            DocumentParseResult with extracted content

        Raises:
            FileNotFoundError: If file does not exist
            UnsupportedFormatError: If format is not supported
            FileTooLargeError: If file exceeds size limit
            DocumentCorruptedError: If document is corrupted
        """
        path = Path(file_path)
        filename = path.name

        logger.debug(
            "Document parsing started (file)",
            extra={
                "file_path": str(path)[:100],
                "document_filename": self._sanitize_filename(filename),
                "project_id": project_id,
            },
        )

        if not path.exists():
            logger.error(
                "Document file not found",
                extra={
                    "file_path": str(path)[:100],
                    "project_id": project_id,
                },
            )
            raise FileNotFoundError(f"File not found: {path}")

        # Read file and parse
        file_bytes = path.read_bytes()
        return self.parse_bytes(file_bytes, filename, project_id)

    def get_supported_formats(self) -> list[str]:
        """Get list of supported file extensions.

        Returns:
            List of supported extensions (e.g., ['.pdf', '.docx', '.txt'])
        """
        return list(SUPPORTED_EXTENSIONS)

    def is_supported(self, filename: str) -> bool:
        """Check if a file format is supported.

        Args:
            filename: Name of the file to check

        Returns:
            True if format is supported, False otherwise
        """
        extension = Path(filename).suffix.lower()
        return extension in SUPPORTED_EXTENSIONS


# Global singleton instance
_document_parser: DocumentParser | None = None


def get_document_parser() -> DocumentParser:
    """Get the global document parser instance.

    Usage:
        from app.utils.document_parser import get_document_parser
        parser = get_document_parser()
        result = parser.parse_bytes(file_bytes, "brand_guide.pdf")
    """
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
        logger.info("DocumentParser singleton created")
    return _document_parser


def parse_document(
    file_bytes: bytes,
    filename: str,
    project_id: str | None = None,
) -> DocumentParseResult:
    """Convenience function for parsing documents.

    Args:
        file_bytes: Document content as bytes
        filename: Name of the file
        project_id: Project ID for logging

    Returns:
        DocumentParseResult with extracted content
    """
    parser = get_document_parser()
    return parser.parse_bytes(file_bytes, filename, project_id)


def parse_document_file(
    file_path: str | Path,
    project_id: str | None = None,
) -> DocumentParseResult:
    """Convenience function for parsing documents from file path.

    Args:
        file_path: Path to the document file
        project_id: Project ID for logging

    Returns:
        DocumentParseResult with extracted content
    """
    parser = get_document_parser()
    return parser.parse_file(file_path, project_id)
